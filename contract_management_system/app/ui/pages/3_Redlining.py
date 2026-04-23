from __future__ import annotations

import io
import re
import sys
from pathlib import Path

import pandas as pd
import streamlit as st
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

sys.path.insert(0, str(Path(__file__).parents[3]))
from app.ui.components.data_loader import (
    get_all_runs,
    get_doc_label_map,
    load_contract_text,
    load_redlines,
    load_risk_signals,
)

st.set_page_config(page_title="Redlining", layout="wide")

st.markdown(
    """
<style>
@import url('https://fonts.googleapis.com/css2?family=Nunito+Sans:wght@400;600;700&family=Varela+Round&display=swap');
html, body, [class*="css"]  { font-family: 'Nunito Sans', sans-serif; }
h1, h2, h3 { font-family: 'Varela Round', sans-serif; }
.card{background:#FFFFFF;border-radius:10px;box-shadow:0 2px 8px rgba(99,102,241,0.08);padding:16px;}
.badge{padding:2px 10px;border-radius:999px;font-size:12px;font-weight:700;border:1px solid transparent;display:inline-block;}
.badge-High{background:#FEF2F2;color:#DC2626;border-color:#FECACA;}
.badge-Medium{background:#FFFBEB;color:#D97706;border-color:#FDE68A;}
.badge-Low{background:#F0FDF4;color:#16A34A;border-color:#BBF7D0;}
</style>
""",
    unsafe_allow_html=True,
)

st.title("Redlining")

runs = get_all_runs()
if not runs:
    st.error("No runs found. Run the pipeline first.")
    st.stop()

selected_run = st.sidebar.selectbox("Run", runs, format_func=lambda p: p.name)
run_key = str(selected_run)

doc_label_map = get_doc_label_map(run_key)
label_to_doc = {v: k for k, v in doc_label_map.items()}

redlines_df = load_redlines(run_key)
risk_df = load_risk_signals(run_key)

if not redlines_df.empty and not risk_df.empty:
    sev_map = risk_df.set_index("rule_id")["severity"].to_dict()
    redlines_df["severity"] = redlines_df["risk_id"].map(sev_map).fillna("Low")

all_labels = list(doc_label_map.values()) if doc_label_map else []
if not all_labels:
    all_labels = redlines_df["document_id"].unique().tolist() if not redlines_df.empty else []

if not all_labels:
    st.info("No contracts are available for redlining.")
    st.stop()

selected_label = st.sidebar.selectbox("Contract", sorted(all_labels))
selected_doc = label_to_doc.get(selected_label, selected_label)

doc_redlines = redlines_df[redlines_df["document_id"] == selected_doc] if not redlines_df.empty else redlines_df
st.markdown(f"### {selected_label}")

contract_text = load_contract_text(run_key, selected_doc)

if not doc_redlines.empty:
    doc_redlines = doc_redlines.reset_index(drop=True)

if contract_text:
    processed = contract_text
    additions = []
    for idx, row in doc_redlines.iterrows():
        orig = str(row.get("original_text") or "")
        prop = str(row.get("proposed_text") or "")
        if orig.strip() and orig in processed:
            repl = (
                f"<del style='background:#FEE2E2;text-decoration:line-through;color:#991B1B'>{orig}</del>"
                f"<ins style='background:#DCFCE7;text-decoration:underline;color:#166534'>{prop}</ins>"
            )
            processed = processed.replace(orig, repl, 1)
        elif not orig.strip() and prop.strip():
            marker = f"<sup style='color:#16A34A'>[+{idx+1}]</sup>"
            processed += marker
            additions.append((idx + 1, prop))

    st.markdown(
        f"""
        <div style="background:#fff;border-radius:10px;padding:24px 32px;
                    box-shadow:0 2px 8px rgba(99,102,241,0.08);
                    max-height:600px;overflow-y:auto;font-family:'Nunito Sans';
                    line-height:1.7;white-space:pre-wrap;">{processed}</div>
        """,
        unsafe_allow_html=True,
    )
    if additions:
        st.markdown("**Additions**")
        for num, txt in additions:
            st.markdown(f"- `[+{num}]` {txt}")
else:
    st.info("Contract text unavailable. Showing fallback card-per-change view.")
    for _, row in doc_redlines.iterrows():
        sev = row.get("severity", "Low")
        st.markdown(f"<div class='card'><b>{row.get('risk_id')}</b> <span class='badge badge-{sev}'>{sev}</span><br><del>{row.get('original_text','')}</del><br><ins>{row.get('proposed_text','')}</ins></div>", unsafe_allow_html=True)

st.divider()
st.subheader("Review Changes")

if doc_redlines.empty:
    st.info("No AI-suggested redlines for this contract.")

with st.expander("Review Changes", expanded=not doc_redlines.empty):
    for i, row in doc_redlines.iterrows():
        risk_id = row.get("risk_id", f"risk_{i}")
        decision_key = f"rl_decision_{selected_doc}_{risk_id}"
        edited_text_key = f"rl_edited_{selected_doc}_{risk_id}"
        sev = row.get("severity", "Low")
        cur = st.session_state.get(decision_key, "Pending")
        c1, c2, c3, c4 = st.columns([2, 2, 6, 4])
        c1.markdown(f"`{risk_id}`")
        c2.markdown(f"<span class='badge badge-{sev}'>{sev}</span>", unsafe_allow_html=True)
        c3.write(row.get("proposed_text", ""))
        with c4:
            b1, b2 = st.columns(2)
            if b1.button("Accept", key=f"accept_{selected_doc}_{i}"):
                st.session_state[decision_key] = "Accepted"
            if b2.button("Reject", key=f"reject_{selected_doc}_{i}"):
                st.session_state[decision_key] = "Rejected"
            edited = st.text_input("Edit", value=st.session_state.get(edited_text_key, row.get("proposed_text", "")), key=f"edit_{selected_doc}_{i}")
            if edited != row.get("proposed_text", ""):
                st.session_state[decision_key] = "Edited"
                st.session_state[edited_text_key] = edited
            st.caption(f"Status: {st.session_state.get(decision_key, cur)}")

def collect_decisions() -> list[dict]:
    out = []
    for _, row in doc_redlines.iterrows():
        risk_id = row.get("risk_id")
        decision_key = f"rl_decision_{selected_doc}_{risk_id}"
        edited_text_key = f"rl_edited_{selected_doc}_{risk_id}"
        decision = st.session_state.get(decision_key, "Pending")
        final_text = row.get("proposed_text", "")
        if decision == "Edited":
            final_text = st.session_state.get(edited_text_key, final_text)
        elif decision == "Rejected":
            final_text = None
        out.append(
            {
                "Contract": selected_label,
                "Risk ID": risk_id,
                "Severity": row.get("severity", "Low"),
                "Original": row.get("original_text", ""),
                "Proposed": row.get("proposed_text", ""),
                "Decision": decision,
                "Final": final_text,
            }
        )
    for c in st.session_state.get(f"rl_custom_{selected_doc}", []):
        out.append(
            {
                "Contract": selected_label,
                "Risk ID": "CUSTOM",
                "Severity": "N/A",
                "Original": "",
                "Proposed": c["proposed_text"],
                "Decision": "Accepted",
                "Final": c["proposed_text"],
                "Clause Ref": c.get("clause_ref", ""),
            }
        )
    return out

st.subheader("Add Custom Change")
custom_key = f"rl_custom_{selected_doc}"
if custom_key not in st.session_state:
    st.session_state[custom_key] = []
with st.form(key=f"custom_form_{selected_doc}", clear_on_submit=True):
    clause_ref = st.text_input("Clause / Section Reference")
    custom_text = st.text_area("Proposed Change Text", height=90)
    submitted = st.form_submit_button("Add Change")
    if submitted and custom_text.strip():
        st.session_state[custom_key].append({"clause_ref": clause_ref, "proposed_text": custom_text})

rows = collect_decisions()

excel_buffer = io.BytesIO()
with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
    pd.DataFrame(rows)[["Contract", "Risk ID", "Severity", "Original", "Proposed", "Decision"]].to_excel(
        writer, sheet_name="Redlines", index=False
    )
    custom_rows = [{"Clause Ref": c.get("clause_ref", ""), "Proposed Text": c.get("proposed_text", "")} for c in st.session_state.get(custom_key, [])]
    pd.DataFrame(custom_rows).to_excel(writer, sheet_name="Custom Changes", index=False)
excel_buffer.seek(0)

doc_buffer = io.BytesIO()
doc = Document()
doc.add_heading(selected_label, level=1)
for r in rows:
    doc.add_heading(f"{r['Risk ID']} ({r['Severity']})", level=2)
    p = doc.add_paragraph()
    run1 = p.add_run(str(r.get("Original", "")))
    run1.font.strike = True
    p.add_run(" → ")
    run2 = p.add_run(str(r.get("Final") or r.get("Proposed", "")))
    run2.font.underline = True
    doc.add_paragraph(f"Decision: {r['Decision']}")
doc.save(doc_buffer)
doc_buffer.seek(0)

pdf_buffer = io.BytesIO()
pdf_doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
styles = getSampleStyleSheet()
story = [Paragraph(f"<b>{selected_label}</b>", styles["BodyText"]), Spacer(1, 10)]
for r in rows:
    story.append(Paragraph(f"<b>{r['Risk ID']} ({r['Severity']})</b>", styles["BodyText"]))
    orig = re.sub(r"[<>&]", "", str(r.get("Original", "")))
    fin = re.sub(r"[<>&]", "", str(r.get("Final") or r.get("Proposed", "")))
    story.append(Paragraph(f"<strike>{orig}</strike> <u>{fin}</u>", styles["BodyText"]))
    story.append(Paragraph(f"Decision: {r['Decision']}", styles["BodyText"]))
    story.append(Spacer(1, 8))
pdf_doc.build(story)
pdf_buffer.seek(0)

st.subheader("Export")
ex1, ex2, ex3, ex4 = st.columns(4)
with ex1:
    st.download_button(
        "Download Excel",
        data=excel_buffer,
        file_name=f"redlines_{selected_doc}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        disabled=not rows,
    )
with ex2:
    st.download_button(
        "Download Word",
        data=doc_buffer,
        file_name=f"redlines_{selected_doc}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        disabled=not rows,
    )
with ex3:
    st.download_button(
        "Download PDF",
        data=pdf_buffer,
        file_name=f"redlines_{selected_doc}.pdf",
        mime="application/pdf",
        disabled=not rows,
    )
with ex4:
    st.button("Send to Coupa", disabled=True, help="Integration coming in Phase 2")
